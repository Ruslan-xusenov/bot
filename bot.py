import asyncio
import re
import logging
import json
import subprocess
import os
from directions_data import directions
from test_data import test_questions
from datetime import datetime
from aiogram import Bot, Dispatcher, Router, types
from aiogram.client.default import DefaultBotProperties
from aiogram.enums import ParseMode
from aiogram.filters import Command, StateFilter
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.types import Message, FSInputFile, CallbackQuery, BufferedInputFile
from docxtpl import DocxTemplate
from aiogram.types import InlineKeyboardMarkup, InlineKeyboardButton
import gspread
from aiogram import F
from aiogram.types import ReplyKeyboardMarkup, KeyboardButton
from docx import Document
from aiogram import types


API_TOKEN = "7636319540:AAHp8fslAIYuiB8JOMoizdt2D4BuQEcJatQ"
ADMIN_ID = 6204272431

bot = Bot(token=API_TOKEN, default=DefaultBotProperties(parse_mode=ParseMode.HTML))
storage = MemoryStorage()
dp = Dispatcher(storage=storage)
router = Router()
dp.include_router(router)

class Form(StatesGroup):
    full_name = State()
    phone = State()
    passport = State()
    passport_given_by = State()
    jshshir = State()
    address = State()
    shape = State()
    direction = State()
    passport_photo = State()
    attestat_photo = State()
    test_index = State()
    test_score = State()

class AdminStates(StatesGroup):
    add_channel = State()
    del_channel = State()
    broadcast = State()
    waiting_for_user_id = State()

def load_channels():
    try:
        with open("channels.json", "r", encoding='utf-8') as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return ["@orientalqabul", "@Oregin_Consult"]

def save_channels(channels):
    with open("channels.json", "w", encoding='utf-8') as f:
        json.dump(channels, f, ensure_ascii=False, indent=2)

REQUIRED_CHANNELS = load_channels()


from aiogram import Bot
import asyncio

async def delete_webhook():
    bot = Bot(token="7636319540:AAHp8fslAIYuiB8JOMoizdt2D4BuQEcJatQ")
    await bot.delete_webhook()
    await bot.session.close()

asyncio.run(delete_webhook())


async def check_subs(user_id):
    for channel in REQUIRED_CHANNELS:
        try:
            member = await bot.get_chat_member(channel, user_id)
            if member.status not in ['creator', 'administrator', 'member']:
                return False
        except Exception as e:
            logging.error(f"Kanal tekshiruv xatosi {channel} uchun: {e}")
            return False
    return True

def log_user(user_id, full_name):
    try:
        with open("user_data.json", "r", encoding='utf-8') as f:
            data = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        data = {}

    data[str(user_id)] = {
        "full_name": full_name,
        "joined": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    }

    with open("user_data.json", "w", encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def log_full_user_data(user_id: int, state_data: dict):
    try:
        with open("user_full_data.json", "r", encoding='utf-8') as f:
            data = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        data = {}

    file_data = {}
    for key in ['passport_photo', 'attestat_photo']:
        if key in state_data:
            file_data[key] = state_data[key]
            del state_data[key]

    data[str(user_id)] = {
        **state_data,
        "joined": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        **file_data
    }

    with open("user_full_data.json", "w", encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

async def check_admin(user_id: int) -> bool:
    return user_id == ADMIN_ID

async def admin_back_button(message: Message):
    reply_keyboard = ReplyKeyboardMarkup(
        keyboard=[
            [KeyboardButton(text="➕ Kanal qo'shish")],
            [KeyboardButton(text="➖ Kanal o'chirish")],
            [KeyboardButton(text="📢 Xabar yuborish")],
            [KeyboardButton(text="📋 Kanal ro'yxati")],
            [KeyboardButton(text="👤 Foydalanuvchi ma'lumotlari")],
            [KeyboardButton(text="Foydalanuvchilar ro'yxati")],
            [KeyboardButton(text="🔙 Bosh menyu")]
        ],
        resize_keyboard=True
    )
    await message.answer("Admin panelga qaytdingiz", reply_markup=reply_keyboard)

def save_users_to_word(users_data: dict, file_name="foydalanuvchilar.docx"):
    if os.path.exists(file_name):
        doc = Document(file_name)
    else:
        doc = Document()
        doc.add_heading("Foydalanuvchilar ro'yxati", 0)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = "Ism"
        hdr[1].text = "Username"
        hdr[2].text = "ID"
        hdr[3].text = "Telefon"
        hdr[4].text = "Pasport"
        hdr[5].text = "JShShIR"

    table = doc.tables[-1]

    for user_id, user in users_data.items():
        if not isinstance(user, dict):
            continue

        row = table.add_row().cells
        row[0].text = user.get("full_name", "")
        row[1].text = f"@{user.get('username', '')}" if user.get("username") else "Yo‘q"
        row[2].text = str(user_id)
        row[3].text = user.get("phone_number", "")
        row[4].text = user.get("passport_number", "")
        row[5].text = user.get("jshshir", "")

    doc.save(file_name)
    return file_name

@router.callback_query(F.data == "save_users_word")
async def save_users_word(callback: CallbackQuery):
    try:
        with open("user_full_data.json", "r", encoding="utf-8") as f:
            users = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        await callback.answer("❌ Foydalanuvchilar ma'lumotlari topilmadi!", show_alert=True)
        return

    file_name = save_users_to_word(users)

    with open(file_name, "rb") as doc_file:
        file_data = doc_file.read()
        input_file = BufferedInputFile(file_data, filename=file_name)
        await callback.message.answer_document(input_file, caption="📄 Foydalanuvchilar ro'yxati Word faylga saqlandi.")

    try:
        os.remove(file_name)
    except Exception as e:
        logging.error(f"Failed to remove temporary file: {e}")

async def get_all_users():
    try:
        with open("user_full_data.json", "r", encoding="utf-8") as f:
            data = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return {}
    users = {}
    for user_id, user_data in data.items():
        users[user_id] = {
            "full_name": user_data.get("full_name", ""),
            "username": user_data.get("username", ""),
            "phone_number": user_data.get("phone", ""),
            "passport_number": user_data.get("passport", ""),
            "jshshir": user_data.get("jshshir", "")
        }
    return users


@router.message(Command("admin"))
async def admin_panel(message: Message):
    if not await check_admin(message.from_user.id):
        return await message.answer("❌ Siz admin emassiz.")

    try:
        with open("user_data.json", "r", encoding='utf-8') as f:
            data = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        data = {}

    today = datetime.now()
    count_1 = count_7 = count_30 = 0

    for user in data.values():
        try:
            join_date = datetime.strptime(user["joined"], "%Y-%m-%d %H:%M:%S")
            days_diff = (today - join_date).days
            if days_diff < 1: count_1 += 1
            if days_diff < 7: count_7 += 1
            if days_diff < 30: count_30 += 1
        except:
            continue

    text = f"""🛠 <b>Admin panel</b>
👤 Bugun qo'shilganlar: <b>{count_1}</b>
📆 7 kunda: <b>{count_7}</b>
📅 30 kunda: <b>{count_30}</b>
📊 Umumiy: <b>{len(data)}</b>
"""
    reply_keyboard = ReplyKeyboardMarkup(
        keyboard=[
            [KeyboardButton(text="➕ Kanal qo'shish")],
            [KeyboardButton(text="➖ Kanal o'chirish")],
            [KeyboardButton(text="📢 Xabar yuborish")],
            [KeyboardButton(text="📋 Kanal ro'yxati")],
            [KeyboardButton(text="👤 Foydalanuvchi ma'lumotlari")],
            [KeyboardButton(text="Foydalanuvchilar ro'yxati")],
            [KeyboardButton(text="🔙 Bosh menyu")],
            [KeyboardButton(text="📄 Word faylga saqlash")]
        ],
        resize_keyboard=True
    )
    await message.answer(text, reply_markup=reply_keyboard)

@router.message(F.text == "➕ Kanal qo'shish")
async def add_channel_handler(message: Message, state: FSMContext):
    if not await check_admin(message.from_user.id):
        return
    
    reply_keyboard = ReplyKeyboardMarkup(
        keyboard=[[KeyboardButton(text="🔙 Orqaga")]],
        resize_keyboard=True
    )
    await message.answer("➕ Kanal username'ini yuboring (masalan: @channel_username):", 
                        reply_markup=reply_keyboard)
    await state.set_state(AdminStates.add_channel)

@router.message(F.text == "➖ Kanal o'chirish")
async def del_channel_handler(message: Message, state: FSMContext):
    if not await check_admin(message.from_user.id):
        return
    
    channels = "\n".join([f"{i+1}. {ch}" for i, ch in enumerate(REQUIRED_CHANNELS)])
    
    reply_keyboard = ReplyKeyboardMarkup(
        keyboard=[[KeyboardButton(text="🔙 Orqaga")]],
        resize_keyboard=True
    )
    
    await message.answer(f"Quyidagi kanallar mavjud:\n{channels}\n\n➖ O'chirmoqchi bo'lgan kanal raqamini yuboring:",
                        reply_markup=reply_keyboard)
    await state.set_state(AdminStates.del_channel)

@router.message(F.text == "📢 Xabar yuborish")
async def broadcast_handler(message: Message, state: FSMContext):
    if not await check_admin(message.from_user.id):
        return
    
    reply_keyboard = ReplyKeyboardMarkup(
        keyboard=[[KeyboardButton(text="🔙 Orqaga")]],
        resize_keyboard=True
    )
    
    await message.answer("✉️ Barcha foydalanuvchilarga yubormoqchi bo'lgan xabaringizni kiriting:",
                        reply_markup=reply_keyboard)
    await state.set_state(AdminStates.broadcast)

@router.message(F.text == "📄 Word faylga saqlash")
async def export_users_to_word(message: Message):
    if not await check_admin(message.from_user.id):
        return await message.answer("❌ Siz admin emassiz.")
    
    try:
        with open("user_full_data.json", "r", encoding="utf-8") as f:
            users = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return await message.answer("❌ Ma'lumotlar topilmadi yoki xatolik yuz berdi.")
    
    if not users:
        return await message.answer("ℹ️ Saqlash uchun foydalanuvchi yo'q.")
    
    file_name = save_users_to_word(users)

    with open(file_name, "rb") as doc_file:
        file_data = doc_file.read()
        input_file = BufferedInputFile(file_data, filename=file_name)
        await message.answer_document(input_file, caption="📄 Foydalanuvchilar ro'yxati hujjatga saqlandi.")

    try:
        os.remove(file_name)
    except Exception as e:
        logging.error(f"Failed to remove temporary file: {e}")

@router.message(F.text == "📋 Kanal ro'yxati")
async def list_channels_handler(message: Message):
    if not await check_admin(message.from_user.id):
        return
    
    channels = "\n".join([f"{i+1}. {ch}" for i, ch in enumerate(REQUIRED_CHANNELS)])
    
    reply_keyboard = ReplyKeyboardMarkup(
        keyboard=[[KeyboardButton(text="🔙 Orqaga")]],
        resize_keyboard=True
    )
    
    await message.answer(f"📋 Majburiy kanallar ro'yxati:\n\n{channels}",
                       reply_markup=reply_keyboard)

@router.message(F.text == "👤 Foydalanuvchi ma'lumotlari")
async def user_data_handler(message: Message, state: FSMContext):
    if not await check_admin(message.from_user.id):
        return
    
    try:
        with open("user_full_data.json", "r", encoding='utf-8') as f:
            data = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return await message.answer("❌ Foydalanuvchi ma'lumotlari topilmadi.")
    
    users = []
    for user_id, user_data in data.items():
        users.append(f"👤 <b>{user_id}</b> - {user_data.get('full_name', 'Nomaʼlum')}")
    
    text = "Foydalanuvchilar ro'yxati:\n\n" + "\n".join(users[:50])
    if len(users) > 50:
        text += f"\n\n... va yana {len(users)-50} ta foydalanuvchi"
    
    reply_keyboard = ReplyKeyboardMarkup(
        keyboard=[[KeyboardButton(text="🔙 Orqaga")]],
        resize_keyboard=True
    )
    
    await message.answer(text, reply_markup=reply_keyboard)
    await message.answer("Maʼlumotlarini koʻrmoqchi boʻlgan foydalanuvchi ID sini yuboring:",
                        reply_markup=reply_keyboard)
    await state.set_state(AdminStates.waiting_for_user_id)

@router.message(F.text == "Foydalanuvchilar ro'yxati")
async def list_all_users(message: Message):
    try:
        with open("user_full_data.json", "r", encoding="utf-8") as f:
            users = json.load(f)
    except Exception:
        await message.answer("❌ Foydalanuvchilar ro'yxatini o'qib bo'lmadi.")
        return

    if not users:
        await message.answer("ℹ️ Hozircha hech qanday foydalanuvchi ro'yxatdan o'tmagan.")
        return

    user_blocks = []
    for user_id, info in users.items():
        lines = [f"🆔 <code>{user_id}</code>"]
        for key, value in info.items():
            if value and key not in ['passport_photo', 'attestat_photo']:
                label = key.replace("_", " ").capitalize()
                lines.append(f"• <b>{label}:</b> {value}")
        user_blocks.append("\n".join(lines))

    text_blocks = []
    current_block = ""
    for block in user_blocks:
        if len(current_block) + len(block) > 4000:
            text_blocks.append(current_block)
            current_block = ""
        current_block += block + "\n\n"
    if current_block:
        text_blocks.append(current_block)

    reply_keyboard = ReplyKeyboardMarkup(
        keyboard=[[KeyboardButton(text="🔙 Orqaga")]],
        resize_keyboard=True
    )

    for block in text_blocks:
        await message.answer(block.strip(), parse_mode="HTML", reply_markup=reply_keyboard)

@router.message(F.text == "🔙 Bosh menyu")
async def back_to_main_menu(message: Message, state: FSMContext):
    await state.clear()
    await message.answer("Bosh menyuga qaytdingiz", reply_markup=types.ReplyKeyboardRemove())
    await message.answer("/start ni bosing")

@router.message(F.text == "🔙 Orqaga")
async def handle_back_button(message: Message, state: FSMContext):
    current_state = await state.get_state()
    
    if current_state in [AdminStates.add_channel, AdminStates.del_channel, 
                        AdminStates.broadcast, AdminStates.waiting_for_user_id]:
        await admin_back_button(message)
        await state.clear()
    elif current_state is None:
        await admin_back_button(message)
    else:
        await message.answer("Bosh menyuga qaytdingiz", reply_markup=types.ReplyKeyboardRemove())

@router.message(AdminStates.add_channel)
async def process_add_channel(message: Message, state: FSMContext):
    if not await check_admin(message.from_user.id):
        return
    
    channel = message.text.strip()
    if channel == "🔙 Orqaga":
        await admin_back_button(message)
        await state.clear()
        return
    
    if not channel.startswith('@'):
        await message.answer("❌ Kanal username'i @ bilan boshlanishi kerak.")
        return
    
    if channel in REQUIRED_CHANNELS:
        await message.answer("❗ Bu kanal allaqachon mavjud.")
    else:
        REQUIRED_CHANNELS.append(channel)
        save_channels(REQUIRED_CHANNELS)
        await message.answer(f"✅ Kanal qo'shildi: {channel}")
    
    await state.clear()
    await admin_panel(message)

@router.message(AdminStates.del_channel)
async def process_del_channel(message: Message, state: FSMContext):
    if not await check_admin(message.from_user.id):
        return
    
    if message.text == "🔙 Orqaga":
        await admin_back_button(message)
        await state.clear()
        return
    
    try:
        index = int(message.text.strip()) - 1
        if 0 <= index < len(REQUIRED_CHANNELS):
            removed_channel = REQUIRED_CHANNELS.pop(index)
            save_channels(REQUIRED_CHANNELS)
            await message.answer(f"❌ Kanal o'chirildi: {removed_channel}")
        else:
            await message.answer("❌ Noto'g'ri raqam kiritildi.")
    except ValueError:
        await message.answer("❌ Faqat raqam kiriting.")
    
    await state.clear()
    await admin_panel(message)

@router.message(AdminStates.broadcast)
async def forward_broadcast(message: Message, state: FSMContext):
    if not await check_admin(message.from_user.id):
        return

    if message.text == "🔙 Orqaga":
        await admin_back_button(message)
        await state.clear()
        return

    try:
        with open("user_data.json", "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception as e:
        return await message.answer("❌ Foydalanuvchilarni o'qib bo'lmadi.")

    success = 0
    errors = 0

    for user_id in data:
        try:
            await bot.forward_message(
                chat_id=user_id,
                from_chat_id=message.chat.id,
                message_id=message.message_id
            )
            success += 1
            await asyncio.sleep(0.1)
        except Exception as e:
            errors += 1
            continue

    await message.answer(f"✅ Yuborildi: {success} ta\n❌ Xato: {errors} ta")
    await state.clear()
    await admin_panel(message)

@router.message(AdminStates.waiting_for_user_id)
async def show_user_data(message: Message, state: FSMContext):
    if not await check_admin(message.from_user.id):
        return
    
    if message.text == "🔙 Orqaga":
        await admin_back_button(message)
        await state.clear()
        return
    
    try:
        with open("user_full_data.json", "r", encoding='utf-8') as f:
            data = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return await message.answer("❌ Foydalanuvchi ma'lumotlari topilmadi.")
    
    user_id = message.text
    if user_id not in data:
        return await message.answer("❌ Foydalanuvchi topilmadi.")
    
    user_data = data[user_id]
    text = f"👤 Foydalanuvchi ma'lumotlari (ID: {user_id}):\n\n"
    
    for key, value in user_data.items():
        if key in ['passport_photo', 'attestat_photo']:
            continue
        text += f"<b>{key.replace('_', ' ').title()}:</b> {value}\n"
    
    await message.answer(text)
    
    if 'passport_photo' in user_data:
        try:
            await message.answer_photo(user_data['passport_photo'], caption="Pasport fotosurati")
        except Exception as e:
            logging.error(f"Error sending passport photo: {e}")
    
    if 'attestat_photo' in user_data:
        try:
            await message.answer_photo(user_data['attestat_photo'], caption="Attestat fotosurati")
        except Exception as e:
            logging.error(f"Error sending attestat photo: {e}")
    
    await state.clear()
    await admin_panel(message)


@router.message(Command("start"))
async def start_handler(message: Message, state: FSMContext):
    log_user(message.from_user.id, message.from_user.full_name)
    
    if not await check_subs(message.from_user.id):
        text = "❗ Botdan foydalanish uchun quyidagi kanallarga obuna bo'ling:\n"
        for ch in REQUIRED_CHANNELS:
            text += f"🔸 <a href='https://t.me/{ch[1:]}'>{ch}</a>\n"
        text += "\n✅ Obuna bo'lib bo'lgach, /start ni qayta yuboring."
        
        keyboard = []
        for ch in REQUIRED_CHANNELS:
            keyboard.append([InlineKeyboardButton(text=f"Obuna bo'lish: {ch}", url=f"https://t.me/{ch[1:]}")])
        keyboard.append([InlineKeyboardButton(text="✅ Tekshirish", callback_data="check_subs")])
        
        markup = InlineKeyboardMarkup(inline_keyboard=keyboard)
        await message.answer(text, reply_markup=markup)
        return

    await message.answer("Assalomu alaykum! Ismingizni to'liq kiriting(Abduqodir Husanov Otabek o'g'li, Abdulla Toirov Abdug'aniyevich):")
    await state.set_state(Form.full_name)

@router.callback_query(lambda c: c.data == "check_subs")
async def process_check_subs(callback_query: types.CallbackQuery, state: FSMContext):
    if await check_subs(callback_query.from_user.id):
        await callback_query.message.delete()
        await callback_query.message.answer("✅ Hammadan obuna bo'ldingiz. Ismingizni to'liq kiriting:")
        await state.set_state(Form.full_name)
    else:
        await callback_query.answer("❌ Hali barcha kanallarga obuna bo'lmagansiz!", show_alert=True)

@router.message(StateFilter(Form.full_name))
async def process_full_name(message: Message, state: FSMContext):
    if len(message.text.split()) < 2:
        await message.answer("❗ Iltimos, ism va familiyangizni to'liq kiriting.")
        return
    
    await state.update_data(full_name=message.text)
    await message.answer("📱 Telefon raqamingizni yuboring (+998911234567):")
    await state.set_state(Form.phone)

@router.message(StateFilter(Form.phone))
async def process_phone(message: Message, state: FSMContext):
    phone = message.text.strip()
    if not re.fullmatch(r"\+998\d{9}", phone):
        await message.answer("❌ Telefon raqam noto'g'ri. +998 bilan boshlanib, 9 ta raqamdan iborat bo'lishi kerak.")
        return
    
    await state.update_data(phone=phone)
    await message.answer("🆔 Pasport seriya va raqamingizni yozing (masalan: AB1234567):")
    await state.set_state(Form.passport)

@router.message(StateFilter(Form.passport))
async def process_passport(message: Message, state: FSMContext):
    passport = message.text.strip().upper()
    if not re.fullmatch(r"[A-Z]{2}\d{7}", passport):
        await message.answer("❌ Noto'g'ri format. Masalan: AB1234567.")
        return
    
    await state.update_data(passport=passport)
    await message.answer("🏛 Pasport qayerdan olingan?(Samarqand viloyati Urgut tumani):")
    await state.set_state(Form.passport_given_by)

@router.message(StateFilter(Form.passport_given_by))
async def process_passport_given_by(message: Message, state: FSMContext):
    await state.update_data(passport_given_by=message.text)
    await message.answer("🔢 JSHSHIR raqamingizni yuboring (12345678910112):")
    await state.set_state(Form.jshshir)

@router.message(StateFilter(Form.jshshir))
async def process_jshshir(message: Message, state: FSMContext):
    jshshir = message.text.strip()
    if not jshshir.isdigit() or len(jshshir) != 14:
        await message.answer("❌ JSHSHIR noto'g'ri. 14 ta raqamdan iborat bo'lishi kerak.")
        return
    
    await state.update_data(jshshir=jshshir)
    await message.answer("🏠 Yashash manzilingizni yozing(Samarqand viloyati Urgut tumani):")
    await state.set_state(Form.address)

@router.message(StateFilter(Form.address))
async def process_address(message: Message, state: FSMContext):
    await state.update_data(address=message.text)
    
    markup = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="Kunduzgi", callback_data="shape_Kunduzgi")],
        [InlineKeyboardButton(text="Kechki", callback_data="shape_Kechki")]
    ])
    await message.answer("🎓 Ta'lim shaklini tanlang:", reply_markup=markup)
    await state.set_state(Form.shape)

@router.callback_query(lambda c: c.data.startswith("shape_"), StateFilter(Form.shape))
async def process_shape(callback_query: types.CallbackQuery, state: FSMContext):
    shape = callback_query.data.split("_")[1]
    await state.update_data(shape=shape)

    dir_names = list(directions.keys())
    keyboard = []
    for i in range(0, len(dir_names), 2):
        row = []
        if i < len(dir_names):
            row.append(KeyboardButton(text=dir_names[i]))
        if i+1 < len(dir_names):
            row.append(KeyboardButton(text=dir_names[i+1]))
        keyboard.append(row)

    reply_markup = ReplyKeyboardMarkup(keyboard=keyboard, resize_keyboard=True)

    await callback_query.message.answer("🎯 Ta'lim yo'nalishini tanlang:", reply_markup=reply_markup)
    await state.set_state(Form.direction)

@router.message(StateFilter(Form.direction))
async def process_direction_text(message: Message, state: FSMContext):
    direction = message.text.strip()
    if direction not in directions:
        await message.answer("❌ Noto‘g‘ri yo‘nalish tanlandi. Iltimos, tugmalardan birini tanlang.")
        return

    data = await state.get_data()
    shape = data.get("shape", "")
    price = directions.get(direction, {}).get(shape, "0")

    await state.update_data(direction=direction)
    await message.answer(
        f"✅ Tanlangan yo'nalish: <b>{direction}</b>\n"
        f"📊 Ta'lim shakli: <b>{shape}</b>\n"
        f"💵 To'lov miqdori: <b>{price} so'm</b>\n\n"
        "📷 Pasportingizning fotosuratini yuboring:"
    )
    await state.set_state(Form.passport_photo)

@router.message(StateFilter(Form.passport_photo), F.photo)
async def step9(message: Message, state: FSMContext):
    await state.update_data(passport_photo=message.photo[-1].file_id)
    await message.answer("📄 Endi attestat fotosuratini yuboring:")
    await state.set_state(Form.attestat_photo)

@router.message(StateFilter(Form.attestat_photo), F.photo)
async def step10(message: Message, state: FSMContext):
    await state.update_data(attestat_photo=message.photo[-1].file_id, score=0, test_index=0)
    await ask_question(message, state)
    await state.set_state(Form.test_index)


async def ask_question(message: Message, state: FSMContext):
    data = await state.get_data()
    index = data.get("test_index", 0)
    
    if index >= len(test_questions):
        await finish_test(message, state)
        return
    
    question = test_questions[index]
    options = "\n".join([f"{chr(65+i)}) {option}" for i, option in enumerate(question['options'])])
    text = f"❓ Savol {index+1}/{len(test_questions)}:\n{question['question']}\n\n{options}"
    
    markup = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text=option, callback_data=f"ans_{i}")]
        for i, option in enumerate(question['options'])
    ])
    
    await message.answer(text, reply_markup=markup)

@router.callback_query(lambda c: c.data.startswith("ans_"), StateFilter(Form.test_index))
async def process_test_answer(callback_query: types.CallbackQuery, state: FSMContext):
    answer_index = int(callback_query.data.split("_")[1])
    data = await state.get_data()
    current_index = data.get("test_index", 0)
    score = data.get("test_score", 0)
    
    correct_index = test_questions[current_index]['answer']
    
    await state.update_data(test_index=current_index+1, test_score=score)
    await ask_question(callback_query.message, state)

async def finish_test(message: Message, state: FSMContext):
    data = await state.get_data()
    score = data.get("test_score", 0)
    total = len(test_questions)
    
    log_full_user_data(message.from_user.id, data)
    
    await message.answer(f"✅ Test yakunlandi!")
    await generate_contract(message, state)


CONTRACT_FILE = "contract_numbers.json"

def get_next_contract_number():
    start_number = 150
    max_number = 9999

    if not os.path.exists(CONTRACT_FILE):
        with open(CONTRACT_FILE, "w") as f:
            json.dump({"last_number": start_number - 1}, f)

    with open(CONTRACT_FILE, "r") as f:
        data = json.load(f)

    last_number = data.get("last_number", start_number - 1)
    next_number = last_number + 1

    if next_number > max_number:
        raise Exception("🔴 Barcha shartnoma raqamlari tugagan (UN9999 gacha)!")

    with open(CONTRACT_FILE, "w") as f:
        json.dump({"last_number": next_number}, f)

    return f"UN{next_number:04d}"

async def generate_contract(message: Message, state: FSMContext):
    try:
        data = await state.get_data()
        today = datetime.now()

        contract_number = get_next_contract_number()
        contract_day = str(today.day)
        contract_month = today.strftime("%m")

        direction = data.get("direction", "")
        shape = data.get("shape", "")
        amount = directions.get(direction, {}).get(shape, "0")
        amount_words = f"{amount} so'm"

        context = {
            "contract_number": contract_number,
            "contract_day": contract_day,
            "contract_month": contract_month,
            "full_name": data.get("full_name", ""),
            "phone_number": data.get("phone", ""),
            "passport": data.get("passport", ""),
            "passport_series": data.get("passport", "")[:2],
            "passport_number": data.get("passport", "")[2:],
            "passport_given_by": data.get("passport_given_by", ""),
            "jshshir": data.get("jshshir", ""),
            "address": data.get("address", ""),
            "shape": shape,
            "direction": direction,
            "contract_amount": amount,
            "contract_amount_words": amount_words,
            "test_score": data.get("test_score", 0)
        }

        template = DocxTemplate("shartnoma_shablon.docx")
        template.render(context)

        file_name = f"{contract_number}.docx"
        pdf_file = f"{contract_number}.pdf"

        template.save(file_name)

        try:
            subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", file_name], check=True)
            if os.path.exists(pdf_file):
                await message.answer(f"📄 Shartnoma tayyorlandi.\nFayl nomi: <b>{pdf_file}</b>")
                
                if data.get("passport_photo"):
                    await message.answer_photo(data["passport_photo"], caption="Pasport fotosurati")
                if data.get("attestat_photo"):
                    await message.answer_photo(data["attestat_photo"], caption="Attestat fotosurati")
                
                await message.answer_document(FSInputFile(pdf_file))
            else:
                await message.answer("❌ PDF fayl topilmadi.")

        except subprocess.CalledProcessError as e:
            await message.answer(f"❌ PDF konvertatsiyasida xatolik: {e}")

        for f in [file_name, pdf_file]:
            try:
                if os.path.exists(f):
                    os.remove(f)
            except Exception as e:
                logging.error(f"Faylni o'chirishda xato: {e}")

    except Exception as e:
        logging.error(f"Shartnoma yaratishda xatolik: {e}")
        await message.answer("❌ Shartnoma yaratishda xatolik yuz berdi. Iltimos, qayta urinib ko'ring.")


@router.errors()
async def errors_handler(update: types.Update, exception: Exception):
    """
    Handle all errors in the bot
    """
    logging.error(f"Update {update} caused error: {exception}", exc_info=True)
    return True


async def main():
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
        handlers=[
            logging.FileHandler("bot.log"),
            logging.StreamHandler()
        ]
    )
    try:
        await dp.start_polling(bot)
    finally:
        await bot.session.close()

if __name__ == "__main__":
    asyncio.run(main())